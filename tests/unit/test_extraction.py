from docx import Document

from app.services.extraction_service import extract


async def test_extract_utf8_text(tmp_path) -> None:
    path = tmp_path / "document.txt"
    path.write_text("First paragraph\n\nSecond paragraph", encoding="utf-8")
    pages = await extract(path, "text/plain")
    assert pages[0].text.startswith("First paragraph")
    assert pages[0].page_number is None


async def test_extract_docx(tmp_path) -> None:
    path = tmp_path / "document.docx"
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    document.save(path)
    pages = await extract(
        path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert pages[0].text == "First paragraph\n\nSecond paragraph"
